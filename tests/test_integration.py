"""
tests/test_integration.py — Full-pipeline integration tests for RecruitX

Every test exercises real components (SQLite, FAISS, embeddings, scoring,
signal analysis, CRUD) with only LLM/OpenRouter calls mocked.

Tests
-----
test_full_pipeline_end_to_end
    Real ranker + real scorers → ranked shortlist with correct structure.
test_empty_result_pipeline
    A deliberately non-matching query produces an empty shortlist.
test_recruit_api_endpoint
    POST /api/recruit returns 200 with real orchestrator (only LLM mocked).
test_recruit_api_invalid_jd
    Empty/missing JD returns 422 validation error.
test_resume_upload_and_search
    Upload a .docx resume → candidate stored in DB and discoverable via FAISS.
test_interview_questions_api
    POST /api/interview-questions returns questions for a known candidate.
"""

import json
import os
import tempfile
from unittest.mock import MagicMock
from uuid import uuid4

import pytest
from fastapi.testclient import TestClient

from api.main import app
from agents.jd_analyst import JDAnalysis
from agents.signal_analyzer import SignalAnalyzerAgent
from scoring.scoring_engine import ScoringEngine
from scoring.skill_gap import SkillGapAnalyzer
from agents.orchestrator import RecruitmentOrchestrator

DEFAULT_JD_ANALYSIS = JDAnalysis(
    required_skills=["Python", "FastAPI", "PostgreSQL", "SQL"],
    preferred_skills=["AWS", "Docker"],
    min_experience_years=3.0,
    education_required="Bachelor's in Computer Science",
    seniority_level="Senior",
    role_summary="Senior Python Developer to build scalable applications.",
    search_query="Senior Python Developer FastAPI PostgreSQL AWS Docker",
)

SAMPLE_JD = (
    "We are looking for a Senior Python Developer with 5+ years experience. "
    "Required skills: Python, FastAPI, PostgreSQL, SQL. "
    "Preferred skills: AWS, Docker. "
    "Education: Bachelor's in Computer Science."
)



# ===================================================================
# Full pipeline — orchestration tests
# ===================================================================


class TestOrchestrationPipeline:
    """Real RecruitmentOrchestrator with injected real dependencies."""

    def test_full_pipeline_end_to_end(self, integration_env):
        """Real ranker, scorer, DB → ranked shortlist with valid structure."""
        ranker = integration_env["ranker"]

        mock_jd = MagicMock()
        mock_jd.analyze.return_value = DEFAULT_JD_ANALYSIS

        orchestrator = RecruitmentOrchestrator(
            jd_analyst=mock_jd,
            candidate_ranker=ranker,
            signal_analyzer=SignalAnalyzerAgent(),
            scoring_engine=ScoringEngine(),
            skill_gap_analyzer=SkillGapAnalyzer(),
        )

        result = orchestrator.run_recruitment_pipeline(SAMPLE_JD, top_k=5)

        # Top-level
        assert result["success"] is True
        assert len(result["shortlist"]) == 5
        assert "jd_analysis" in result
        assert isinstance(result["processing_time_ms"], float)

        # JD analysis
        jd = result["jd_analysis"]
        assert "Python" in jd["required_skills"]
        assert jd["min_experience_years"] == 3.0
        assert jd["seniority_level"] == "Senior"

        # First entry
        entry = result["shortlist"][0]
        assert entry["rank"] == 1
        assert entry["candidate_id"] > 0
        assert 0 <= entry["final_score"] <= 100
        assert 0 <= entry["semantic_score"] <= 100
        assert 0 <= entry["skill_score"] <= 100
        assert 0 <= entry["signal_score"] <= 100
        assert isinstance(entry["explanation"], str) and len(entry["explanation"]) > 0
        assert entry["explanation"].startswith("Rank #1.")

        # Skill gap
        sg = entry["skill_gap"]
        assert "matched" in sg and "missing" in sg and "bonus" in sg
        assert len(sg["matched"]) > 0  # at least some skills matched

        # Ranking order
        for i in range(len(result["shortlist"]) - 1):
            assert result["shortlist"][i]["final_score"] >= result["shortlist"][i + 1]["final_score"]

    def test_unusual_query_returns_low_scores(self, integration_env):
        """Non-matching query returns candidates with low semantic scores (~50)."""
        ranker = integration_env["ranker"]

        mock_jd = MagicMock()
        mock_jd.analyze.return_value = JDAnalysis(
            required_skills=["Golang", "Rust"],
            preferred_skills=["WebAssembly"],
            min_experience_years=5.0,
            education_required="Bachelor's",
            seniority_level="Senior",
            role_summary="Systems Engineer",
            search_query="Golang Rust WebAssembly systems programming low latency",
        )

        orchestrator = RecruitmentOrchestrator(
            jd_analyst=mock_jd,
            candidate_ranker=ranker,
        )

        result = orchestrator.run_recruitment_pipeline(
            "Golang and Rust systems engineer with WebAssembly.",
            top_k=10,
        )

        assert result["success"] is True
        assert len(result["shortlist"]) == 10
        assert "jd_analysis" in result
        assert result["processing_time_ms"] >= 0
        # FAISS always returns nearest neighbors; scores will be low (~50)
        for entry in result["shortlist"]:
            assert entry["semantic_score"] < 70


# ===================================================================
# API endpoint — real components, LLM mocked
# ===================================================================


class TestRecruitAPI:
    """POST /api/recruit with real orchestrator (LLM mocked via monkeypatch)."""

    @pytest.fixture(autouse=True)
    def _mock_jd_llm(self, monkeypatch):
        """Replace JD Analyst's LLM call with a canned analysis."""
        monkeypatch.setattr(
            "agents.jd_analyst.JDAnalystAgent.analyze",
            lambda self, jd: DEFAULT_JD_ANALYSIS,
        )

    def test_successful_recruit(self):
        """Valid JD returns 200 with ranked shortlist."""
        client = TestClient(app)
        response = client.post("/api/recruit", json={
            "job_description": SAMPLE_JD,
            "top_k": 3,
        })

        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert len(data["shortlist"]) == 3
        assert data["shortlist"][0]["rank"] == 1
        assert "required_skills" in data["jd_analysis"]
        assert data["processing_time_ms"] > 0

    def test_empty_jd_returns_422(self):
        """Empty job description triggers Pydantic validation → 422."""
        client = TestClient(app)
        response = client.post("/api/recruit", json={"job_description": ""})
        assert response.status_code == 422

    def test_missing_jd_returns_422(self):
        """Missing job_description field triggers Pydantic validation → 422."""
        client = TestClient(app)
        response = client.post("/api/recruit", json={})
        assert response.status_code == 422


# ===================================================================
# Resume upload → candidate searchable via FAISS
# ===================================================================


class TestResumeUploadIntegration:
    """Upload a .docx resume, verify it is stored and discoverable."""

    @pytest.fixture(autouse=True)
    def _mock_llm_parse(self, monkeypatch):
        """Replace the LLM parse chain with a canned response.

        Only the LLM is mocked; text extraction, hashing, DB save,
        and FAISS update run with real components.
        """
        mock_chain = MagicMock()
        mock_chain.invoke.return_value = json.dumps({
            "name": "Integration Tester",
            "email": "integration.test@recruitx.io",
            "phone": "9999999999",
            "location": "Pune",
            "skills": ["Python", "FastAPI", "PostgreSQL", "Docker", "AWS"],
            "experience_years": 5.0,
            "education": "B.Tech Computer Science",
            "previous_roles": ["Senior Backend Dev at TestCorp"],
        })
        monkeypatch.setattr(
            "utils.resume_parser.ResumeParser._create_parse_chain",
            lambda self: mock_chain,
        )

    def _create_test_docx(self, path: str) -> str:
        """Create a minimal .docx file with candidate info and return its path."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Integration Tester")
        doc.add_paragraph("Email: integration.test@recruitx.io")
        doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL, Docker, AWS")
        doc.add_paragraph("Experience: 5 years")
        doc.add_paragraph("Education: B.Tech Computer Science")
        doc.add_paragraph("Previous Roles: Senior Backend Dev at TestCorp")
        doc.save(path)
        return path

    def test_upload_and_search(self, integration_env):
        """Upload .docx → candidate in DB → discoverable via recruit API."""
        # Create a temporary .docx file
        fd, docx_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            self._create_test_docx(docx_path)

            with open(docx_path, "rb") as f:
                docx_bytes = f.read()

            client = TestClient(app)

            # --- Upload ---
            upload_resp = client.post(
                "/api/upload-resume",
                files={
                    "file": (
                        "integration_resume.docx",
                        docx_bytes,
                        "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document",
                    )
                },
            )

            assert upload_resp.status_code == 200, upload_resp.text
            upload_data = upload_resp.json()
            assert upload_data["is_new"] is True
            assert upload_data["candidate"]["name"] == "Integration Tester"
            new_id = upload_data["candidate"]["id"]

            # --- Verify candidate in DB via API ---
            get_resp = client.get("/api/candidates")
            assert get_resp.status_code == 200
            all_emails = [c["email"] for c in get_resp.json()["candidates"]]
            assert "integration.test@recruitx.io" in all_emails

            # --- Verify candidate is discoverable via FAISS ---
            import agents.jd_analyst as jd_mod

            original_analyze = jd_mod.JDAnalystAgent.analyze
            jd_mod.JDAnalystAgent.analyze = lambda self, jd: DEFAULT_JD_ANALYSIS

            try:
                recruit_resp = client.post("/api/recruit", json={
                    "job_description": (
                        "Senior Python Developer with FastAPI, PostgreSQL, "
                        "Docker and AWS experience."
                    ),
                    "top_k": 50,
                })
            finally:
                jd_mod.JDAnalystAgent.analyze = original_analyze

            assert recruit_resp.status_code == 200
            candidate_ids = [e["candidate_id"] for e in recruit_resp.json()["shortlist"]]
            assert new_id in candidate_ids, (
                f"Candidate {new_id} not found in search results"
            )

        finally:
            if os.path.exists(docx_path):
                os.remove(docx_path)

    def _create_test_docx_with_email(self, path: str, email: str, name: str = None) -> str:
        """Create a minimal .docx with a specific email address."""
        from docx import Document

        doc = Document()
        doc.add_paragraph(name or "Integration Tester")
        doc.add_paragraph(f"Email: {email}")
        doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL, Docker, AWS")
        doc.add_paragraph("Experience: 5 years")
        doc.add_paragraph("Education: B.Tech Computer Science")
        doc.add_paragraph("Previous Roles: Senior Backend Dev at TestCorp")
        doc.save(path)
        return path

    def test_duplicate_email_returns_409(self, integration_env):
        """Upload .docx once, then upload a different file with the same email → 409."""
        unique_email = f"dup.{uuid4().hex}@recruitx.io"

        fd1, docx_path1 = tempfile.mkstemp(suffix=".docx")
        os.close(fd1)
        try:
            self._create_test_docx_with_email(docx_path1, unique_email)

            with open(docx_path1, "rb") as f:
                docx_bytes1 = f.read()

            # Override LLM mock for the first doc (same canned profile but different email)
            mock_chain1 = MagicMock()
            mock_chain1.invoke.return_value = json.dumps({
                "name": "Duplicate Test User",
                "email": unique_email,
                "phone": "8888888888",
                "location": "Mumbai",
                "skills": ["Python", "FastAPI"],
                "experience_years": 3.0,
                "education": "B.Sc",
                "previous_roles": ["Dev"],
            })
            import utils.resume_parser as rp
            original_create_chain = rp.ResumeParser._create_parse_chain

            client = TestClient(app)

            # First upload — succeeds
            rp.ResumeParser._create_parse_chain = lambda self: mock_chain1
            try:
                resp1 = client.post(
                    "/api/upload-resume",
                    files={
                        "file": (
                            "first_resume.docx",
                            docx_bytes1,
                            "application/vnd.openxmlformats-officedocument"
                            ".wordprocessingml.document",
                        )
                    },
                )
            finally:
                rp.ResumeParser._create_parse_chain = original_create_chain

            assert resp1.status_code == 200, resp1.text
            assert resp1.json()["is_new"] is True

            # Second upload — same email, different content → 409
            fd2, docx_path2 = tempfile.mkstemp(suffix=".docx")
            os.close(fd2)
            try:
                self._create_test_docx_with_email(
                    docx_path2, unique_email, name="Duplicate Test User v2",
                )

                with open(docx_path2, "rb") as f:
                    docx_bytes2 = f.read()

                mock_chain2 = MagicMock()
                mock_chain2.invoke.return_value = json.dumps({
                    "name": "Duplicate Test User v2",
                    "email": unique_email,
                    "phone": "8888888888",
                    "location": "Mumbai",
                    "skills": ["Python", "Go", "Kubernetes"],
                    "experience_years": 6.0,
                    "education": "M.Sc",
                    "previous_roles": ["Senior Dev"],
                })

                rp.ResumeParser._create_parse_chain = lambda self: mock_chain2
                try:
                    resp2 = client.post(
                        "/api/upload-resume",
                        files={
                            "file": (
                                "second_resume.docx",
                                docx_bytes2,
                                "application/vnd.openxmlformats-officedocument"
                                ".wordprocessingml.document",
                            )
                        },
                    )
                finally:
                    rp.ResumeParser._create_parse_chain = original_create_chain

                assert resp2.status_code == 409, resp2.text
                assert "already exists" in resp2.json()["detail"]
            finally:
                if os.path.exists(docx_path2):
                    os.remove(docx_path2)

        finally:
            if os.path.exists(docx_path1):
                os.remove(docx_path1)


# ===================================================================
# Interview Question Generator API
# ===================================================================


class TestInterviewQuestionsAPI:
    """POST /api/interview-questions with mocked LLM."""

    @pytest.fixture(autouse=True)
    def _mock_generator_llm(self, monkeypatch):
        """Replace LLM generation with canned questions."""
        from utils.interview_generator import InterviewQuestions

        canned = InterviewQuestions(questions=[
            "How would you design a FastAPI application for high concurrency?",
            "Describe your experience with PostgreSQL query optimization.",
            "You have strong Python skills. Walk me through a complex project.",
            "The role requires Docker. How would you learn it?",
            "Tell me about a time you handled a production outage.",
        ])

        monkeypatch.setattr(
            "api.routes.interviews._generator",
            MagicMock(generate=MagicMock(return_value=canned)),
        )

    def test_generate_questions_for_known_candidate(self):
        """Returns questions with correct structure."""
        client = TestClient(app)
        response = client.post("/api/interview-questions", json={
            "candidate_id": 1,
            "candidate_name": "Rahul Sharma",
            "candidate_skills": "Python, FastAPI, SQL",
            "candidate_experience_years": 5.0,
            "candidate_education": "B.Tech CS",
            "candidate_previous_roles": "Senior Dev at TCS",
            "job_description": "Senior Python Developer with FastAPI experience.",
            "skill_gap_matched": ["Python", "FastAPI"],
            "skill_gap_missing": ["Docker"],
            "skill_gap_bonus": ["SQL"],
            "explanation": "Rank #1. Matched 2 skill(s), missing 1 skill(s).",
        })

        assert response.status_code == 200
        data = response.json()
        assert data["candidate_id"] == 1
        assert len(data["questions"]) == 5
        assert all(isinstance(q, str) and len(q) > 0 for q in data["questions"])

    def test_missing_fields_returns_422(self):
        """Request without required fields returns 422."""
        client = TestClient(app)
        response = client.post("/api/interview-questions", json={})
        assert response.status_code == 422
